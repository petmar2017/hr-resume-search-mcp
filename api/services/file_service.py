"""File service for handling document processing and text extraction."""

import os
import uuid
import aiofiles
from datetime import datetime, timedelta
from pathlib import Path
from typing import Optional, Dict, Any
import pypdf
import docx
import magic


class FileService:
    """Service for file operations and text extraction."""
    
    def __init__(self):
        """Initialize file service."""
        self.mime = magic.Magic(mime=True)
    
    async def extract_text(self, file_path: Path, file_extension: str) -> str:
        """
        Extract text from various file formats.
        
        Args:
            file_path: Path to the file
            file_extension: File extension (e.g., '.pdf', '.docx')
            
        Returns:
            Extracted text content
        """
        file_extension = file_extension.lower()
        
        if file_extension == ".pdf":
            return self._extract_pdf_text(file_path)
        elif file_extension in [".doc", ".docx"]:
            return self._extract_docx_text(file_path)
        else:
            # Try to read as plain text
            return self._extract_plain_text(file_path)
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        text = ""
        try:
            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
        except Exception as e:
            raise Exception(f"Failed to extract text from PDF: {str(e)}")
        
        return text.strip()
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        text = ""
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"Failed to extract text from DOCX: {str(e)}")
        
        return text.strip()
    
    def _extract_plain_text(self, file_path: Path) -> str:
        """
        Extract text from plain text file.
        
        Args:
            file_path: Path to text file
            
        Returns:
            File content
        """
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as file:
                return file.read()
        except Exception as e:
            raise Exception(f"Failed to read text file: {str(e)}")
    
    def validate_file(self, file_path: Path, expected_type: Optional[str] = None) -> bool:
        """
        Validate file type and integrity.
        
        Args:
            file_path: Path to file
            expected_type: Expected MIME type
            
        Returns:
            True if file is valid
        """
        if not file_path.exists():
            return False
        
        if expected_type:
            detected_type = self.mime.from_file(str(file_path))
            return detected_type == expected_type
        
        return True
    
    def get_file_metadata(self, file_path: Path) -> dict:
        """
        Get file metadata.
        
        Args:
            file_path: Path to file
            
        Returns:
            File metadata dictionary
        """
        if not file_path.exists():
            return {}
        
        stats = file_path.stat()
        
        return {
            "size": stats.st_size,
            "created": stats.st_ctime,
            "modified": stats.st_mtime,
            "mime_type": self.mime.from_file(str(file_path)),
            "extension": file_path.suffix
        }
    
    def sanitize_filename(self, filename: str) -> str:
        """
        Sanitize filename for safe storage.
        
        Args:
            filename: Original filename
            
        Returns:
            Sanitized filename
        """
        # Remove path components
        filename = os.path.basename(filename)
        
        # Replace unsafe characters
        unsafe_chars = ["<", ">", ":", '"', "/", "\\", "|", "?", "*"]
        for char in unsafe_chars:
            filename = filename.replace(char, "_")
        
        # Limit length
        name, ext = os.path.splitext(filename)
        if len(name) > 100:
            name = name[:100]
        
        return name + ext
    
    # Additional methods expected by tests
    def validate_file_type(self, filename: str) -> bool:
        """
        Validate file type based on extension.
        
        Args:
            filename: Name of the file to validate
            
        Returns:
            True if file type is allowed
        """
        allowed_extensions = [".pdf", ".doc", ".docx"]
        file_path = Path(filename)
        return file_path.suffix.lower() in allowed_extensions
    
    def validate_file_size(self, size_bytes: int) -> bool:
        """
        Validate file size is within allowed limits.
        
        Args:
            size_bytes: Size of file in bytes
            
        Returns:
            True if file size is valid
        """
        max_size = 10 * 1024 * 1024  # 10MB
        return 0 < size_bytes <= max_size
    
    async def save_file(self, filename: str, content: bytes) -> str:
        """
        Save file content to disk.
        
        Args:
            filename: Original filename
            content: File content as bytes
            
        Returns:
            Path to saved file
        """
        # Create uploads directory if it doesn't exist
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        
        # Generate unique filename
        unique_filename = self.generate_unique_filename(filename)
        file_path = upload_dir / unique_filename
        
        # Save file
        async with aiofiles.open(file_path, "wb") as f:
            await f.write(content)
        
        return str(file_path)
    
    async def extract_text_from_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF file (async wrapper).
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        return self._extract_pdf_text(Path(file_path))
    
    async def extract_text_from_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file (async wrapper).
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        return self._extract_docx_text(Path(file_path))
    
    async def process_uploaded_file(self, filename: str, content: bytes) -> Dict[str, Any]:
        """
        Process uploaded file: save, validate, and extract text.
        
        Args:
            filename: Original filename
            content: File content
            
        Returns:
            Processing results including file path and extracted text
        """
        # Validate file type
        if not self.validate_file_type(filename):
            raise ValueError(f"Unsupported file type: {filename}")
        
        # Validate file size
        if not self.validate_file_size(len(content)):
            raise ValueError(f"File size too large: {len(content)} bytes")
        
        # Save file
        file_path = await self.save_file(filename, content)
        
        # Extract text based on file extension
        file_extension = Path(filename).suffix.lower()
        if file_extension == ".pdf":
            text = await self.extract_text_from_pdf(file_path)
        elif file_extension in [".doc", ".docx"]:
            text = await self.extract_text_from_docx(file_path)
        else:
            text = ""
        
        # Get file metadata
        metadata = self.get_file_metadata(Path(file_path))
        
        return {
            "file_path": file_path,
            "original_filename": filename,
            "extracted_text": text,
            "metadata": metadata,
            "processed_at": datetime.utcnow().isoformat()
        }
    
    def generate_unique_filename(self, original_filename: str) -> str:
        """
        Generate unique filename while preserving extension.
        
        Args:
            original_filename: Original filename
            
        Returns:
            Unique filename
        """
        # Sanitize original filename
        sanitized = self.sanitize_filename(original_filename)
        
        # Split name and extension
        name, ext = os.path.splitext(sanitized)
        
        # Generate unique identifier
        unique_id = str(uuid.uuid4())[:8]
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        
        # Create unique filename
        unique_name = f"{name}_{timestamp}_{unique_id}{ext}"
        
        return unique_name
    
    async def delete_file(self, file_path: str) -> None:
        """
        Delete file from disk.
        
        Args:
            file_path: Path to file to delete
        """
        path = Path(file_path)
        if path.exists():
            path.unlink()
    
    async def cleanup_old_files(self, days_old: int = 30) -> None:
        """
        Clean up files older than specified days.
        
        Args:
            days_old: Number of days after which files should be deleted
        """
        upload_dir = Path("uploads")
        if not upload_dir.exists():
            return
        
        cutoff_date = datetime.utcnow() - timedelta(days=days_old)
        
        for file_path in upload_dir.glob("*"):
            if file_path.is_file():
                # Get file modification time
                file_mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
                
                # Delete if older than cutoff
                if file_mtime < cutoff_date:
                    file_path.unlink()